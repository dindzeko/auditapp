import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import io


# =========================================================
# KONSTANTA (BARU)
# =========================================================
# STRUCTURE_TOLERANCE : dipakai untuk MENDETEKSI struktur tabel
#   (apakah suatu baris adalah subtotal / total implisit).
#   Versi lama: max(5, abs(nilai) * 0.00001) -> proporsional, membengkak
#   jadi ratusan ribu untuk nilai miliaran. Diganti absolut kecil.
#
# DEFAULT_FOOTING_TOLERANCE : toleransi VERIFIKASI footing (Rp).
#   Bisa diatur dari UI. Dulu proporsional (bug: makin besar angka,
#   makin longgar). Sekarang absolut & dikontrol auditor.
STRUCTURE_TOLERANCE = 2.0
DEFAULT_FOOTING_TOLERANCE = 1.0


# =========================================================
# STREAMLIT APP
# =========================================================

def app():
    # =====================================================
    # PERUBAHAN #1: st.set_page_config() DIHAPUS.
    # Fungsi ini dipanggil sebagai sub-halaman oleh router app.py,
    # sedangkan set_page_config hanya boleh dipanggil SEKALI dan HARUS
    # jadi perintah Streamlit pertama. Kalau dibiarkan -> error saat
    # RecalTab dibuka lewat menu AuditApp.
    # Untuk layout wide seluruh app, taruh
    #     st.set_page_config(page_title="AuditApp", layout="wide")
    # sebagai baris pertama di app.py.
    # =====================================================

    st.title("📄 Verifikasi Footing dan Persentase Tabel Word")

    st.write(
        """
        Upload dokumen Wajib Bentuk Word `.docx`. Aplikasi akan membantu melakukan telstruk.
        """
    )

    st.info(
        """
        Tanda hasil pemeriksaan:
        - `^` warna hijau = sesuai / verified.
        - `X` warna merah = berbeda dengan hasil hitung sistem.
        """
    )

    st.caption(
        "DISCLAIMER : PERHITUNGAN SISTEM BISA JADI SALAH. APABILA DITEMUKAN ANGKA SELISIH MAKA PERIKSA KEMBALI DAN KONFIRMASI DENGAN ENTITAS BILA DIPERLUKAN"
    )

    col1, col2 = st.columns(2)

    with col1:
        tambah_baris_rekalkulasi = st.checkbox(
            "Jika tidak ada baris JUMLAH/TOTAL/total implisit, tambahkan baris Rekalkulasi Sistem",
            value=False,
            key="recaltab_tambah_baris"
        )

    with col2:
        cek_persentase = st.checkbox(
            "Cek kolom persentase otomatis",
            value=True,
            key="recaltab_cek_persen"
        )

    # PERUBAHAN #2: kontrol toleransi footing (absolut, default 1 Rp).
    col3, col4 = st.columns(2)

    with col3:
        footing_tolerance = st.number_input(
            "Toleransi footing (Rp)",
            min_value=0.0,
            value=DEFAULT_FOOTING_TOLERANCE,
            step=1.0,
            help=(
                "Selisih maksimal antara total tertulis dan hasil hitung sistem "
                "agar tetap dianggap SESUAI. Untuk pembulatan biasanya cukup 1-2 rupiah. "
                "JANGAN dibuat besar; toleransi longgar bisa menutupi selisih footing."
            ),
            key="recaltab_footing_tol"
        )

    with col4:
        tampilkan_debug = st.checkbox(
            "Tampilkan ringkasan proses",
            value=False,
            key="recaltab_debug"
        )

    uploaded_file = st.file_uploader(
        "Upload File Word (.docx)",
        type=["docx"],
        key="recaltab_uploader"
    )

    if uploaded_file is not None:
        try:
            doc = Document(uploaded_file)

            with st.spinner("Memproses dokumen..."):
                summary = recalculate_tables(
                    doc=doc,
                    tambah_baris_rekalkulasi=tambah_baris_rekalkulasi,
                    cek_persentase=cek_persentase,
                    footing_tolerance=footing_tolerance
                )

                output = io.BytesIO()
                doc.save(output)
                output.seek(0)

            st.success("Rekalkulasi selesai!")

            # PERUBAHAN #3: peringatan merged cell (sel gabungan).
            tabel_merged = summary.get("tabel_dengan_merged_cell", [])

            if tabel_merged:
                st.warning(
                    "⚠️ Terdeteksi sel gabungan (merged cell) pada tabel nomor: "
                    f"{', '.join(map(str, tabel_merged))}. "
                    "Pada tabel dengan merged cell, pemetaan kolom bisa bergeser sehingga "
                    "hasil footing berpotensi TIDAK akurat. Sistem sudah mencoba mencegah "
                    "penghitungan ganda, namun hasil pada tabel ini WAJIB direviu manual."
                )

            if tampilkan_debug:
                st.subheader("Ringkasan Proses")
                st.json(summary)

            nama_file_hasil = buat_nama_file_hasil(uploaded_file.name)

            st.download_button(
                label="📥 Unduh Hasil Rekalkulasi",
                data=output,
                file_name=nama_file_hasil,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="recaltab_download"
            )

        except Exception as e:
            st.error(f"Terjadi kesalahan: {str(e)}")
            st.error("Pastikan file yang diupload adalah dokumen Word `.docx` valid.")


def buat_nama_file_hasil(nama_file_upload):
    if not nama_file_upload:
        return "hasil_Rekalkulasi.docx"

    if nama_file_upload.lower().endswith(".docx"):
        nama_file_tanpa_ext = nama_file_upload[:-5]
    else:
        nama_file_tanpa_ext = nama_file_upload

    return f"{nama_file_tanpa_ext}_Rekalkulasi.docx"


# =========================================================
# HELPER MERGED CELL (BARU - PERUBAHAN #3)
# =========================================================

def row_has_horizontal_merge(row):
    """
    True jika ada dua sel dalam satu baris yang merujuk ke elemen <w:tc>
    yang sama (indikasi horizontal merge). python-docx mengembalikan objek
    _Cell yang sama berulang kali untuk sel gabungan.
    """
    seen = set()

    for cell in row.cells:
        tc_id = id(cell._tc)

        if tc_id in seen:
            return True

        seen.add(tc_id)

    return False


def table_has_merged_cell(table):
    for row in table.rows:
        if row_has_horizontal_merge(row):
            return True

    return False


# =========================================================
# MAIN PROCESS
# =========================================================

def recalculate_tables(doc, tambah_baris_rekalkulasi=False, cek_persentase=True,
                       footing_tolerance=DEFAULT_FOOTING_TOLERANCE):
    summary = {
        "jumlah_tabel": len(doc.tables),
        "tabel_diproses": 0,
        "tabel_tanpa_kolom_numerik": 0,
        "tabel_dengan_merged_cell": [],   # PERUBAHAN #3
        "baris_total_ditemukan": 0,
        "baris_total_implisit_ditemukan": 0,
        "baris_subtotal_dilewati": 0,
        "sel_footing_verified": 0,
        "sel_footing_berbeda": 0,
        "baris_rekalkulasi_ditambahkan": 0,
        "kolom_persen_dicek": 0,
        "sel_persen_verified": 0,
        "sel_persen_berbeda": 0,
        "formula_persen_ditemukan": []
    }

    for table_idx, table in enumerate(doc.tables, start=1):
        if not table.rows:
            continue

        # PERUBAHAN #3: catat tabel yang punya merged cell.
        if table_has_merged_cell(table):
            summary["tabel_dengan_merged_cell"].append(table_idx)

        clean_table_old_marks(table)

        numeric_cols = detect_numeric_columns_for_footing(table)

        if not numeric_cols:
            summary["tabel_tanpa_kolom_numerik"] += 1

            if cek_persentase:
                verify_percentage_columns(
                    table=table,
                    summary=summary,
                    table_idx=table_idx
                )

            continue

        summary["tabel_diproses"] += 1

        total_indices = find_total_row_indices(table)

        implicit_total_indices = []

        if not total_indices:
            implicit_total_indices = find_implicit_total_row_indices(
                table=table,
                numeric_cols=numeric_cols
            )

        final_total_indices = total_indices if total_indices else implicit_total_indices

        if final_total_indices:
            if total_indices:
                summary["baris_total_ditemukan"] += len(total_indices)
            else:
                summary["baris_total_implisit_ditemukan"] += len(implicit_total_indices)

            final_total_idx = final_total_indices[-1]

            vertical_sums, skipped_subtotal_count = calculate_sums_before_total_row(
                table=table,
                total_row_idx=final_total_idx,
                numeric_cols=numeric_cols
            )

            summary["baris_subtotal_dilewati"] += skipped_subtotal_count

            total_row = table.rows[final_total_idx]

            result = verify_total_row(
                total_row=total_row,
                numeric_cols=numeric_cols,
                vertical_sums=vertical_sums,
                tolerance=footing_tolerance
            )

            summary["sel_footing_verified"] += result["verified"]
            summary["sel_footing_berbeda"] += result["different"]

        else:
            if tambah_baris_rekalkulasi:
                vertical_sums, skipped_subtotal_count = calculate_sums_all_rows(
                    table=table,
                    numeric_cols=numeric_cols
                )

                summary["baris_subtotal_dilewati"] += skipped_subtotal_count

                added = add_recalculation_row(
                    table=table,
                    numeric_cols=numeric_cols,
                    vertical_sums=vertical_sums
                )

                if added:
                    summary["baris_rekalkulasi_ditambahkan"] += 1

        if cek_persentase:
            verify_percentage_columns(
                table=table,
                summary=summary,
                table_idx=table_idx
            )

    return summary


# =========================================================
# SUM CALCULATION
# =========================================================
# PERUBAHAN #3: dedup berdasarkan identitas elemen <w:tc>.
# Pada baris dengan merged cell, satu nilai numerik bisa muncul di
# beberapa indeks kolom -> tanpa dedup akan dijumlahkan ganda.

def _sum_row_into(vertical_sums, row, numeric_cols, seen_tc_ids):
    for col_idx in numeric_cols:
        if col_idx >= len(row.cells):
            continue

        cell = row.cells[col_idx]
        tc_id = id(cell._tc)

        # Jangan hitung sel yang sama dua kali (akibat horizontal merge).
        if tc_id in seen_tc_ids:
            continue

        seen_tc_ids.add(tc_id)

        number = parse_number(cell.text, dash_as_zero=True)

        if number is None:
            continue

        vertical_sums[col_idx] += number


def calculate_sums_before_total_row(table, total_row_idx, numeric_cols):
    vertical_sums = [0.0] * len(table.columns)
    skipped_subtotal_count = 0

    for row_idx in range(0, total_row_idx):
        row = table.rows[row_idx]

        skip, reason = should_skip_row_automatically(
            table=table,
            row=row,
            row_idx=row_idx,
            numeric_cols=numeric_cols
        )

        if skip:
            if reason == "subtotal":
                skipped_subtotal_count += 1
            continue

        seen_tc_ids = set()
        _sum_row_into(vertical_sums, row, numeric_cols, seen_tc_ids)

    return vertical_sums, skipped_subtotal_count


def calculate_sums_all_rows(table, numeric_cols):
    vertical_sums = [0.0] * len(table.columns)
    skipped_subtotal_count = 0

    for row_idx, row in enumerate(table.rows):
        skip, reason = should_skip_row_automatically(
            table=table,
            row=row,
            row_idx=row_idx,
            numeric_cols=numeric_cols
        )

        if skip:
            if reason == "subtotal":
                skipped_subtotal_count += 1
            continue

        seen_tc_ids = set()
        _sum_row_into(vertical_sums, row, numeric_cols, seen_tc_ids)

    return vertical_sums, skipped_subtotal_count


def should_skip_row_automatically(table, row, row_idx, numeric_cols):
    if is_header_number_row(row):
        return True, "header_number"

    if is_total_row(row):
        return True, "total"

    if is_likely_header_text_row(row):
        return True, "header_text"

    if is_probable_subtotal_row(table, row, row_idx, numeric_cols):
        return True, "subtotal"

    return False, ""


# =========================================================
# TOTAL ROW DETECTION
# =========================================================

def find_total_row_indices(table):
    total_indices = []

    for idx, row in enumerate(table.rows):
        if is_total_row(row):
            total_indices.append(idx)

    return total_indices


def is_total_row(row):
    keywords = [
        "JUMLAH",
        "TOTAL",
        "GRANDTOTAL",
        "GRAND TOTAL"
    ]

    non_empty_texts = []

    for cell in row.cells:
        text = normalize_text(cell.text)

        if text:
            non_empty_texts.append(text)

    if not non_empty_texts:
        return False

    first_text = non_empty_texts[0]

    for keyword in keywords:
        key = normalize_text(keyword)

        if first_text == key:
            return True

        if first_text.startswith(key):
            return True

    return False


def find_implicit_total_row_indices(table, numeric_cols):
    """
    Mendeteksi baris total tanpa label.

    Contoh:
    |                    | Rp | 44.670.257.956 |
    atau
    |                    |    | 44.670.257.956 |

    Logika:
    - dicari dari bagian bawah tabel;
    - baris punya angka;
    - uraian kosong / sangat minim / hanya simbol;
    - atau nilai numeriknya bold;
    - atau nilainya cocok dengan hasil penjumlahan baris di atasnya.
    """

    rows = list(table.rows)
    candidates = []

    start_idx = max(0, len(rows) - 8)

    for idx in range(len(rows) - 1, start_idx - 1, -1):
        row = rows[idx]

        if is_header_number_row(row):
            continue

        if is_total_row(row):
            continue

        numeric_values = get_numeric_values(row, numeric_cols)

        if not numeric_values:
            continue

        if is_implicit_total_row_candidate(table, idx, numeric_cols):
            candidates.append(idx)
            break

    return sorted(candidates)


def is_implicit_total_row_candidate(table, row_idx, numeric_cols):
    row = table.rows[row_idx]

    if row_idx <= 0:
        return False

    numeric_values = get_numeric_values(row, numeric_cols)

    if not numeric_values:
        return False

    label_text = get_row_label_text(row, numeric_cols)
    label_norm = normalize_text(label_text)

    label_is_empty_or_weak = (
        label_norm == ""
        or label_norm in ["RP", "-", "–", "—"]
        or len(label_norm) <= 3
    )

    numeric_is_bold = numeric_cells_are_bold(row, numeric_cols)

    vertical_sums, _ = calculate_sums_before_total_row(
        table=table,
        total_row_idx=row_idx,
        numeric_cols=numeric_cols
    )

    match_count = 0
    checked_count = 0

    for col_idx in numeric_cols:
        if col_idx >= len(row.cells):
            continue

        existing_number = parse_number(row.cells[col_idx].text, dash_as_zero=True)

        if existing_number is None:
            continue

        calculated_number = vertical_sums[col_idx]

        if abs(existing_number) < 1:
            continue

        checked_count += 1
        # PERUBAHAN #2: toleransi struktur absolut (bukan proporsional).
        tolerance = STRUCTURE_TOLERANCE

        if numbers_are_equal(existing_number, calculated_number, tolerance):
            match_count += 1

    if checked_count > 0 and match_count >= 1:
        return True

    if label_is_empty_or_weak and numeric_is_bold:
        return True

    if label_is_empty_or_weak and row_idx == len(table.rows) - 1:
        return True

    return False


def get_row_label_text(row, numeric_cols):
    parts = []

    for idx, cell in enumerate(row.cells):
        text = cell.text.strip()

        if idx in numeric_cols:
            continue

        if not text:
            continue

        clean = normalize_text(text)

        if clean in ["RP", "-", "–", "—"]:
            continue

        if parse_number(text, dash_as_zero=False) is not None:
            continue

        parts.append(text)

    return " ".join(parts).strip()


def numeric_cells_are_bold(row, numeric_cols):
    total_runs = 0
    bold_runs = 0

    for col_idx in numeric_cols:
        if col_idx >= len(row.cells):
            continue

        cell = row.cells[col_idx]

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if run.text.strip():
                    total_runs += 1

                    if run.bold:
                        bold_runs += 1

    if total_runs == 0:
        return False

    return bold_runs >= max(1, int(total_runs * 0.5))


# =========================================================
# NUMERIC COLUMN DETECTION
# =========================================================

def detect_numeric_columns_for_footing(table, sample_rows=15):
    numeric_cols = []
    data_rows = []

    for row in table.rows:
        if is_header_number_row(row):
            continue

        if is_total_row(row):
            continue

        if is_likely_header_text_row(row):
            continue

        data_rows.append(row)

        if len(data_rows) >= sample_rows:
            break

    for col_idx in range(len(table.columns)):
        if is_percent_column(table, col_idx):
            continue

        numeric_count = 0

        for row in data_rows:
            if col_idx >= len(row.cells):
                continue

            text = row.cells[col_idx].text.strip()

            if not text:
                continue

            number = parse_number(text, dash_as_zero=True)

            if number is not None:
                numeric_count += 1

        if numeric_count >= 1:
            numeric_cols.append(col_idx)

    return numeric_cols


def detect_all_money_value_columns(table):
    cols = []

    for col_idx in range(len(table.columns)):
        if is_percent_column(table, col_idx):
            continue

        count = 0

        for row in table.rows:
            if is_header_number_row(row):
                continue

            if is_likely_header_text_row(row):
                continue

            if col_idx >= len(row.cells):
                continue

            number = parse_number(row.cells[col_idx].text, dash_as_zero=True)

            if number is not None:
                count += 1

            if count >= 2:
                cols.append(col_idx)
                break

    return cols


# =========================================================
# ROW / COLUMN DETECTION
# =========================================================

def is_header_number_row(row):
    values = []

    for cell in row.cells:
        text = cell.text.strip()
        text = text.replace(" ", "")
        text = text.replace("\n", "")
        text = text.replace("\r", "")

        if text:
            values.append(text)

    if not values:
        return False

    pattern = re.compile(r"^\d+(\s*=\s*[\d\+\-\*/\(\)]+)?$")

    matched = 0

    for value in values:
        if pattern.match(value):
            matched += 1

    return matched >= max(2, int(len(values) * 0.6))


def is_likely_header_text_row(row):
    non_empty = []

    for cell in row.cells:
        text = cell.text.strip()

        if text:
            non_empty.append(text)

    if not non_empty:
        return True

    numeric_found = 0

    for text in non_empty:
        if parse_number(text, dash_as_zero=False) is not None:
            numeric_found += 1

    return numeric_found == 0


def is_percent_column(table, col_idx):
    header_text = ""

    for row in table.rows[:8]:
        if col_idx < len(row.cells):
            header_text += " " + normalize_text(row.cells[col_idx].text)

    # Deteksi utama berbasis header (paling andal).
    if "%" in header_text:
        return True

    if "PERSEN" in header_text:
        return True

    if "PERSENTASE" in header_text:
        return True

    if "RASIO" in header_text:
        return True

    if "NAIK" in header_text and "TURUN" in header_text:
        return True

    # =====================================================
    # PERUBAHAN #5: fallback berbasis rentang nilai DIPERKETAT.
    # Versi lama: -500..10000 dgn min 3 sampel -> terlalu longgar,
    # kolom kuantitas / rupiah kecil ikut dianggap persen lalu
    # DIKECUALIKAN dari cek footing (berbahaya).
    # Sekarang: rentang -100..300, butuh mayoritas kuat (>=80%),
    # DAN minimal ada satu nilai berdesimal (persen umumnya berdesimal;
    # kolom "jumlah unit" biasanya bilangan bulat).
    # Konsekuensi disengaja: bila persen kebetulan semua bulat & tanpa
    # header, kolom tsb tidak auto-terdeteksi. Ini sisi aman -> lebih
    # baik persen tak tercek daripada kolom uang luput dari footing.
    # =====================================================
    sample_values = []

    for row in table.rows:
        if is_header_number_row(row):
            continue

        if is_likely_header_text_row(row):
            continue

        if col_idx >= len(row.cells):
            continue

        raw = row.cells[col_idx].text.strip()

        if raw in ["", "-", "–", "—"]:
            continue

        number = parse_number(raw, dash_as_zero=False)

        if number is not None:
            sample_values.append(number)

        if len(sample_values) >= 8:
            break

    if len(sample_values) >= 4:
        in_range = sum(1 for v in sample_values if -100 <= v <= 300)
        has_decimal = any(abs(v - round(v)) > 1e-9 for v in sample_values)

        threshold = max(4, int(len(sample_values) * 0.8))

        if in_range >= threshold and has_decimal:
            return True

    return False


def detect_percentage_columns(table):
    cols = []

    for col_idx in range(len(table.columns)):
        if is_percent_column(table, col_idx):
            cols.append(col_idx)

    return sorted(list(set(cols)))


# =========================================================
# SUBTOTAL DETECTION
# =========================================================

def is_probable_subtotal_row(table, row, row_idx, numeric_cols):
    current_values = get_numeric_values(row, numeric_cols)

    if not current_values:
        return False

    if not row_has_meaningful_text(row):
        return False

    candidate_children = collect_candidate_child_rows(
        table=table,
        row_idx=row_idx,
        numeric_cols=numeric_cols,
        max_children=20
    )

    if not candidate_children:
        return False

    checked = 0
    matched = 0

    for col_idx in numeric_cols:
        if col_idx >= len(row.cells):
            continue

        current_value = parse_number(row.cells[col_idx].text, dash_as_zero=True)

        if current_value is None or abs(current_value) < 1:
            continue

        child_sum = 0
        child_count = 0

        for child_row in candidate_children:
            if col_idx >= len(child_row.cells):
                continue

            child_value = parse_number(child_row.cells[col_idx].text, dash_as_zero=True)

            if child_value is not None:
                child_sum += child_value
                child_count += 1

        if child_count < 1:
            continue

        checked += 1
        # PERUBAHAN #2: toleransi struktur absolut.
        tolerance = STRUCTURE_TOLERANCE

        if numbers_are_equal(current_value, child_sum, tolerance):
            matched += 1

    if checked == 0:
        return False

    if is_bold_row(row) and matched >= 1:
        return True

    if len(candidate_children) >= 2 and matched >= 1:
        return True

    return False


def collect_candidate_child_rows(table, row_idx, numeric_cols, max_children=20):
    candidate_children = []
    rows = list(table.rows)

    for next_idx in range(row_idx + 1, len(rows)):
        next_row = rows[next_idx]

        if is_header_number_row(next_row):
            continue

        if is_total_row(next_row):
            break

        if is_likely_header_text_row(next_row):
            continue

        if candidate_children and is_bold_row(next_row):
            break

        next_values = get_numeric_values(next_row, numeric_cols)

        if next_values:
            candidate_children.append(next_row)

        if len(candidate_children) >= max_children:
            break

    return candidate_children


def row_has_meaningful_text(row):
    for cell in row.cells:
        text = cell.text.strip()

        if text and parse_number(text, dash_as_zero=False) is None:
            return True

    return False


def get_numeric_values(row, numeric_cols):
    values = {}

    for col_idx in numeric_cols:
        if col_idx >= len(row.cells):
            continue

        number = parse_number(row.cells[col_idx].text, dash_as_zero=True)

        if number is not None:
            values[col_idx] = number

    return values


def is_bold_row(row):
    total_runs = 0
    bold_runs = 0

    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if run.text.strip():
                    total_runs += 1

                    if run.bold:
                        bold_runs += 1

    if total_runs == 0:
        return False

    return bold_runs >= max(1, int(total_runs * 0.5))


# =========================================================
# FOOTING VERIFICATION
# =========================================================

def verify_total_row(total_row, numeric_cols, vertical_sums, tolerance=DEFAULT_FOOTING_TOLERANCE):
    result = {
        "verified": 0,
        "different": 0
    }

    # PERUBAHAN #3: hindari menandai sel yang sama dua kali (merged cell).
    seen_tc_ids = set()

    for col_idx in numeric_cols:
        if col_idx >= len(total_row.cells):
            continue

        cell = total_row.cells[col_idx]

        tc_id = id(cell._tc)
        if tc_id in seen_tc_ids:
            continue
        seen_tc_ids.add(tc_id)

        existing_number = parse_number(cell.text, dash_as_zero=True)

        if existing_number is None:
            continue

        calculated_number = vertical_sums[col_idx]

        # PERUBAHAN #2: toleransi absolut dari UI (bukan proporsional).
        if numbers_are_equal(existing_number, calculated_number, tolerance):
            add_status_mark(cell, "^", RGBColor(0, 176, 80))
            add_recalculation_note_to_cell(
                cell=cell,
                calculated_number=calculated_number,
                is_percent=False,
                color=RGBColor(0, 176, 80)
            )
            result["verified"] += 1
        else:
            add_status_mark(cell, "X", RGBColor(255, 0, 0))
            add_recalculation_note_to_cell(
                cell=cell,
                calculated_number=calculated_number,
                is_percent=False,
                color=RGBColor(255, 0, 0)
            )
            result["different"] += 1

    return result


# =========================================================
# PERCENTAGE VERIFICATION
# =========================================================

def verify_percentage_columns(table, summary=None, table_idx=None):
    percent_cols = detect_percentage_columns(table)
    money_cols = detect_all_money_value_columns(table)

    for percent_col in percent_cols:
        formula = infer_percentage_formula(table, percent_col, money_cols)

        if formula is None:
            continue

        result = apply_percentage_formula_check(table, percent_col, formula)

        if summary is not None:
            summary["kolom_persen_dicek"] += 1
            summary["sel_persen_verified"] += result["verified"]
            summary["sel_persen_berbeda"] += result["different"]
            summary["formula_persen_ditemukan"].append({
                "tabel": table_idx,
                "kolom_persen": percent_col + 1,
                "formula": formula["description"],
                "verified": result["verified"],
                "different": result["different"]
            })


def infer_percentage_formula(table, percent_col, money_cols, min_match=2):
    candidates = build_candidate_percentage_formulas(table, percent_col, money_cols)

    best_formula = None
    best_match = -1
    best_tested = 0

    for formula in candidates:
        tested = 0
        matched = 0

        for row in table.rows:
            if is_header_number_row(row):
                continue

            if is_likely_header_text_row(row):
                continue

            expected = get_cell_number(row, percent_col, dash_as_zero=False)

            if expected is None:
                continue

            calculated = calculate_formula_value(row, formula)

            if calculated is None:
                continue

            tested += 1

            if percentage_numbers_equal(expected, calculated):
                matched += 1

            if tested >= 10:
                break

        if tested == 0:
            continue

        if matched > best_match:
            best_match = matched
            best_tested = tested
            best_formula = formula

    if best_formula is None:
        return None

    if best_match >= min_match:
        return best_formula

    if best_tested > 0 and best_match / best_tested >= 0.70:
        return best_formula

    return None


def build_candidate_percentage_formulas(table, percent_col, money_cols):
    candidates = []
    header_text = get_column_header_text(table, percent_col)

    left_money_cols = [col for col in money_cols if col < percent_col]

    if len(left_money_cols) >= 2:
        numerator = left_money_cols[-1]
        denominator = left_money_cols[-2]

        candidates.append({
            "type": "ratio",
            "numerator_col": numerator,
            "denominator_col": denominator,
            "priority": 0,
            "description": f"Kolom {numerator + 1} / Kolom {denominator + 1} x 100"
        })

    if "NAIK" in header_text or "TURUN" in header_text:
        previous_candidates = [col for col in money_cols if col < percent_col]

        if previous_candidates:
            previous_col = previous_candidates[-1]

            for current_col in money_cols:
                if current_col == previous_col:
                    continue

                if current_col < previous_col:
                    candidates.append({
                        "type": "growth",
                        "current_col": current_col,
                        "previous_col": previous_col,
                        "priority": 0,
                        "description": f"(Kolom {current_col + 1} - Kolom {previous_col + 1}) / Kolom {previous_col + 1} x 100"
                    })

    for numerator_col in money_cols:
        for denominator_col in money_cols:
            if numerator_col == denominator_col:
                continue

            candidates.append({
                "type": "ratio",
                "numerator_col": numerator_col,
                "denominator_col": denominator_col,
                "priority": 10 + abs(percent_col - numerator_col) + abs(percent_col - denominator_col),
                "description": f"Kolom {numerator_col + 1} / Kolom {denominator_col + 1} x 100"
            })

            candidates.append({
                "type": "growth",
                "current_col": numerator_col,
                "previous_col": denominator_col,
                "priority": 20 + abs(percent_col - numerator_col) + abs(percent_col - denominator_col),
                "description": f"(Kolom {numerator_col + 1} - Kolom {denominator_col + 1}) / Kolom {denominator_col + 1} x 100"
            })

    candidates.sort(key=lambda x: x["priority"])

    return candidates


def get_column_header_text(table, col_idx):
    text = ""

    for row in table.rows[:8]:
        if col_idx < len(row.cells):
            text += " " + normalize_text(row.cells[col_idx].text)

    return text


def calculate_formula_value(row, formula):
    if formula["type"] == "ratio":
        numerator = get_cell_number(row, formula["numerator_col"], dash_as_zero=True)
        denominator = get_cell_number(row, formula["denominator_col"], dash_as_zero=True)

        if numerator is None or denominator is None or denominator == 0:
            return None

        return numerator / denominator * 100

    if formula["type"] == "growth":
        current = get_cell_number(row, formula["current_col"], dash_as_zero=True)
        previous = get_cell_number(row, formula["previous_col"], dash_as_zero=True)

        if current is None or previous is None or previous == 0:
            return None

        return (current - previous) / previous * 100

    return None


def apply_percentage_formula_check(table, percent_col, formula):
    result = {
        "verified": 0,
        "different": 0
    }

    for row in table.rows:
        if is_header_number_row(row):
            continue

        if is_likely_header_text_row(row):
            continue

        if percent_col >= len(row.cells):
            continue

        existing_value = get_cell_number(row, percent_col, dash_as_zero=False)

        if existing_value is None:
            continue

        calculated_value = calculate_formula_value(row, formula)

        if calculated_value is None:
            continue

        cell = row.cells[percent_col]

        if percentage_numbers_equal(existing_value, calculated_value):
            add_status_mark(cell, "^", RGBColor(0, 176, 80))
            add_recalculation_note_to_cell(
                cell=cell,
                calculated_number=calculated_value,
                is_percent=True,
                color=RGBColor(0, 176, 80)
            )
            result["verified"] += 1
        else:
            add_status_mark(cell, "X", RGBColor(255, 0, 0))
            add_recalculation_note_to_cell(
                cell=cell,
                calculated_number=calculated_value,
                is_percent=True,
                color=RGBColor(255, 0, 0)
            )
            result["different"] += 1

    return result


def get_cell_number(row, col_idx, dash_as_zero=True):
    if col_idx >= len(row.cells):
        return None

    return parse_number(row.cells[col_idx].text, dash_as_zero=dash_as_zero)


def percentage_numbers_equal(a, b, tolerance=0.05):
    return abs(a - b) <= tolerance


# =========================================================
# MARKING
# =========================================================

def _delete_paragraph(paragraph):
    """Hapus elemen paragraf sepenuhnya (bukan sekadar kosongkan teks)."""
    element = paragraph._element
    parent = element.getparent()

    if parent is not None:
        parent.remove(element)

    paragraph._p = None
    paragraph._element = None


def clean_table_old_marks(table):
    for row in table.rows:
        for cell in row.cells:
            clean_existing_marks_and_notes(cell)


def add_status_mark(cell, mark, color):
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    run = paragraph.add_run(f" {mark}")
    run.font.name = "Calibri"
    run.font.size = Pt(16)   # ukuran 16 dipakai sbagai penanda mark (lihat cleanup)
    run.font.bold = True
    run.font.color.rgb = color


def clean_existing_marks_and_notes(cell):
    # =====================================================
    # PERUBAHAN #4: pembersihan mark yang AMAN.
    # Versi lama melakukan text.replace("X","") / ("^","") pada SEMUA
    # run -> huruf X kapital pada teks sah (TAX, BOX, IX, TXN, dst)
    # ikut terhapus diam-diam saat file diproses ulang.
    #
    # Sekarang:
    # - Paragraf catatan "Rekalkulasi:" dihapus sebagai ELEMEN (tidak
    #   menyisakan paragraf kosong yang menumpuk).
    # - Run mark dihapus HANYA jika isinya persis "^"/"X" DAN berukuran
    #   16pt (ciri khas mark yang dibuat add_status_mark). Teks biasa
    #   tidak tersentuh.
    # =====================================================
    for paragraph in list(cell.paragraphs):
        if "Rekalkulasi:" in paragraph.text:
            _delete_paragraph(paragraph)
            continue

        for run in list(paragraph.runs):
            if run.text.strip() in ["^", "X"] and run.font.size == Pt(16):
                run.text = ""


def add_recalculation_note_to_cell(cell, calculated_number, is_percent=False, color=None):
    paragraph = cell.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    if is_percent:
        text = f"Rekalkulasi: {format_percent(calculated_number)}"
    else:
        text = f"Rekalkulasi: {format_number(calculated_number)}"

    if color is None:
        color = RGBColor(255, 0, 0)

    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = color


# =========================================================
# ADD RECALCULATION ROW
# =========================================================

def add_recalculation_row(table, numeric_cols, vertical_sums):
    if not any(abs(vertical_sums[col]) > 0 for col in numeric_cols):
        return False

    new_row = table.add_row()
    new_row.cells[0].text = "Rekalkulasi Sistem"

    for col_idx in range(len(table.columns)):
        if col_idx in numeric_cols and abs(vertical_sums[col_idx]) > 0:
            cell = new_row.cells[col_idx]
            cell.text = format_number(vertical_sums[col_idx])
            set_recalculation_cell_style(cell)

    return True


def set_recalculation_cell_style(cell):
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.bold = True


# =========================================================
# NUMBER PARSER AND FORMATTER
# =========================================================

def parse_number(text, dash_as_zero=True):
    if text is None:
        return None

    text = str(text).strip()

    if text == "":
        return None

    text = text.replace("\n", "")
    text = text.replace("\r", "")
    text = text.replace("\t", "")
    text = text.replace(" ", "")

    text = text.replace("Rp", "")
    text = text.replace("RP", "")
    text = text.replace("rp", "")
    text = text.replace("%", "")

    text = text.replace("^", "")
    text = text.replace("X", "")

    text = re.sub(r"Rekalkulasi:[-\d\.\,\(\)]+", "", text)

    if text in ["-", "–", "—"]:
        return 0.0 if dash_as_zero else None

    is_negative = False

    if re.match(r"^\(.+\)$", text):
        is_negative = True
        text = text[1:-1]

    text = text.replace(".", "")
    text = text.replace(",", ".")

    if not re.match(r"^-?\d+(\.\d+)?$", text):
        return None

    try:
        number = float(text)

        if is_negative:
            number = -abs(number)

        return number

    except ValueError:
        return None


def format_number(number):
    if number is None:
        return ""

    return f"{number:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def format_percent(number):
    if number is None:
        return ""

    return f"{number:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def numbers_are_equal(a, b, tolerance=1):
    return abs(a - b) <= tolerance


def normalize_text(text):
    if text is None:
        return ""

    return (
        str(text)
        .replace(" ", "")
        .replace("\n", "")
        .replace("\r", "")
        .replace("\t", "")
        .strip()
        .upper()
    )


# =========================================================
# RUN APP
# =========================================================

if __name__ == "__main__":
    app()
